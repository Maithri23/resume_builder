import io, json, re
from xhtml2pdf import pisa
from docx import Document
from bs4 import BeautifulSoup
from flask import render_template

def _nonempty_text(s):
    return bool(s and str(s).strip())

def _nonempty_item(obj):
    if isinstance(obj, str):
        return _nonempty_text(obj)
    if isinstance(obj, dict):
        return any(_nonempty_item(v) for v in obj.values())
    if isinstance(obj, list):
        return any(_nonempty_item(v) for v in obj)
    return bool(obj)

def sanitize_resume_payload(payload):
    if not isinstance(payload, dict):
        return {}
    out = {}
    for k, v in payload.items():
        if k in ("skills", "achievements"):
            cleaned = [s for s in (v or []) if _nonempty_text(s)]
            if cleaned:
                out[k] = cleaned
        elif k in ("experience", "education", "projects", "certifications"):
            cleaned_list = []
            for item in (v or []):
                if isinstance(item, dict):
                    cleaned = {kk: (vv or "").strip() for kk, vv in item.items() if _nonempty_text(vv)}
                    if _nonempty_item(cleaned):
                        cleaned_list.append(cleaned)
            if cleaned_list:
                out[k] = cleaned_list
        elif k in ("summary", ):
            text = (v or {}).get("text", "").strip()
            if text:
                out[k] = {"text": text}
        elif k == "personal":
            # keep only provided fields
            cleaned = {kk: vv.strip() for kk, vv in (v or {}).items() if _nonempty_text(vv)}
            if cleaned:
                out[k] = cleaned
        else:
            if _nonempty_item(v):
                out[k] = v
    return out

def render_resume_html(resume, sanitize=False):
    payload = {}
    try:
        payload = json.loads(resume.data_json or "{}")
    except Exception:
        payload = {}
    if sanitize:
        payload = sanitize_resume_payload(payload)
    return render_template("resume_export.html", r=resume, data=payload)

def html_to_pdf(html: str):
    pdf = io.BytesIO()
    pisa.CreatePDF(io.StringIO(html), dest=pdf)
    pdf.seek(0)
    return pdf

def html_to_docx(html: str):
    soup = BeautifulSoup(html, "html.parser")
    doc = Document()

    # Extract simple blocks
    # Only write headings that have following content in the soup
    for block in soup.find_all(['h1','h2','h3','p','li']):
        text = block.get_text(" ", strip=True)
        if not text:
            continue
        # Skip section headings with no sibling content (heuristic: next tag not li/p)
        if block.name in ('h2','h3'):
            # look ahead for at least one non-empty p/li before next heading
            sib = block.find_next_sibling()
            has_content = False
            while sib and sib.name not in ('h1','h2','h3'):
                if sib.name in ('p','li') and sib.get_text(strip=True):
                    has_content = True
                    break
                sib = sib.find_next_sibling()
            if not has_content:
                continue

        if block.name == 'h1':
            doc.add_heading(text, level=1)
        elif block.name == 'h2':
            doc.add_heading(text, level=2)
        elif block.name == 'h3':
            doc.add_heading(text, level=3)
        elif block.name == 'li':
            p = doc.add_paragraph(style='List Bullet'); p.add_run(text)
        else:
            doc.add_paragraph(text)

    buf = io.BytesIO()
    doc.save(buf); buf.seek(0)
    return buf

def html_to_txt(html: str, title: str):
    soup = BeautifulSoup(html, "html.parser")
    lines = [title, "="*len(title), ""]
    for el in soup.find_all(['h2','p','li']):
        text = el.get_text(" ", strip=True)
        if text: lines.append(text)
    buf = io.BytesIO(("\n".join(lines)).encode("utf-8"))
    buf.seek(0); return buf
